import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

class CustomPDF(FPDF):
    def header(self):
        pass
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def slugify(company, title):
    text = f"{company}_{title}"
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '_', text)
    return text.strip('_').lower()

def clean_non_ascii(text):
    return text.encode('ascii', 'ignore').decode('ascii')

def save_docx_cv(cv_data, filepath):
    doc = Document()
    
    # Page setup (1-inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styling helper
    def add_para(text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(clean_non_ascii(text))
        run.font.size = Pt(size)
        run.bold = bold
        run.font.italic = italic
        run.font.name = 'Arial'
        return p

    # Name and Title
    add_para("Kaushal Sonawane", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(cv_data.get("title", "Software Engineer"), size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para("sonawanekaushal05@gmail.com | +91 8767515720 | LinkedIn | GitHub | Portfolio", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Professional Summary
    add_para("Summary", size=13, bold=True, space_after=4)
    add_para(cv_data.get("summary", ""), size=10.5, space_after=12)

    # Work Experience
    add_para("Work Experience", size=13, bold=True, space_after=6)
    for exp in cv_data.get("experience", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_role = p.add_run(clean_non_ascii(exp.get("role", "")))
        r_role.bold = True
        r_role.font.size = Pt(11)
        
        p_comp = doc.add_paragraph()
        p_comp.paragraph_format.space_after = Pt(4)
        r_comp = p_comp.add_run(clean_non_ascii(f"{exp.get('company', '')} | {exp.get('dates', '')}"))
        r_comp.font.italic = True
        r_comp.font.size = Pt(10)
        
        for bullet in exp.get("bullets", []):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(3)
            r_bullet = p_bullet.add_run(clean_non_ascii(bullet))
            r_bullet.font.size = Pt(10.5)

    # Education
    add_para("Education", size=13, bold=True, space_after=6)
    edu = cv_data.get("education", {})
    if edu:
        p_deg = doc.add_paragraph()
        p_deg.paragraph_format.space_after = Pt(2)
        r_deg = p_deg.add_run(clean_non_ascii(edu.get("degree", "")))
        r_deg.bold = True
        r_deg.font.size = Pt(11)
        
        p_inst = doc.add_paragraph()
        p_inst.paragraph_format.space_after = Pt(4)
        r_inst = p_inst.add_run(clean_non_ascii(f"{edu.get('institution', '')} | {edu.get('dates', '')}"))
        r_inst.font.italic = True
        r_inst.font.size = Pt(10)
        
        if edu.get("details"):
            p_det = doc.add_paragraph()
            p_det.paragraph_format.space_after = Pt(6)
            r_det = p_det.add_run(clean_non_ascii(edu.get("details", "")))
            r_det.font.size = Pt(10.5)
            
    # Skills
    add_para("Skills", size=13, bold=True, space_after=6)
    skills = cv_data.get("skills", {})
    for category, items in skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_cat = p.add_run(clean_non_ascii(f"{category}: "))
        r_cat.bold = True
        r_cat.font.size = Pt(10.5)
        r_items = p.add_run(clean_non_ascii(items))
        r_items.font.size = Pt(10.5)

    # Certifications
    add_para("Certifications", size=13, bold=True, space_after=6)
    for cert in cv_data.get("certifications", []):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_cert = p.add_run(clean_non_ascii(cert))
        r_cert.font.size = Pt(10.5)

    # Projects
    add_para("Projects", size=13, bold=True, space_after=6)
    for proj in cv_data.get("projects", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_name = p.add_run(clean_non_ascii(proj.get("name", "")))
        r_name.bold = True
        r_name.font.size = Pt(11)
        
        for bullet in proj.get("bullets", []):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(3)
            r_bullet = p_bullet.add_run(clean_non_ascii(bullet))
            r_bullet.font.size = Pt(10.5)

    doc.save(filepath)

def save_docx_cover_letter(letter_data, filepath):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    def add_para(text, size=11, bold=False, italic=False, space_after=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(clean_non_ascii(text))
        run.font.size = Pt(size)
        run.bold = bold
        run.font.italic = italic
        run.font.name = 'Arial'
        return p

    # Sender Info
    add_para("Kaushal Sonawane", size=12, bold=True, space_after=2)
    add_para("Pune, Maharashtra, India\nsonawanekaushal05@gmail.com | +91 8767515720", size=10, space_after=18)

    # Date
    import datetime
    current_date = datetime.date.today().strftime("%B %d, %Y")
    add_para(current_date, size=11, space_after=12)

    # Recipient Info
    add_para("Hiring Manager\nTarget Company", size=11, space_after=12)

    # Salutation
    add_para("Dear Hiring Manager,", size=11, space_after=12)

    # Body
    for para in letter_data.get("body_paragraphs", []):
        add_para(para, size=11, space_after=12)

    # Sign-off
    add_para("Sincerely,", size=11, space_after=4)
    add_para("Kaushal Sonawane", size=11, bold=True, space_after=0)

    doc.save(filepath)

def save_pdf_from_markdown(markdown_text, filepath):
    pdf = CustomPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    
    for line in markdown_text.split('\n'):
        line_stripped = line.strip()
        if not line_stripped:
            pdf.ln(2)
            continue
            
        cleaned_line = clean_non_ascii(line_stripped)
        
        # Reset cursor X to left margin (15) to prevent FPDF coordinate calculation errors
        pdf.set_x(15)
            
        if cleaned_line.startswith('### '):
            pdf.set_font("Helvetica", 'B', 11)
            pdf.ln(3)
            pdf.set_x(15)
            pdf.multi_cell(0, 5, cleaned_line.replace('### ', ''))
            pdf.set_font("Helvetica", size=9.5)
        elif cleaned_line.startswith('## '):
            pdf.set_font("Helvetica", 'B', 13)
            pdf.ln(4)
            pdf.set_x(15)
            pdf.multi_cell(0, 6, cleaned_line.replace('## ', ''))
            pdf.set_font("Helvetica", size=9.5)
        elif cleaned_line.startswith('# '):
            pdf.set_font("Helvetica", 'B', 16)
            pdf.ln(5)
            pdf.set_x(15)
            pdf.multi_cell(0, 8, cleaned_line.replace('# ', ''), align='C')
            pdf.set_font("Helvetica", size=9.5)
        elif cleaned_line.startswith('- ') or cleaned_line.startswith('* '):
            bullet_text = cleaned_line[2:].replace('**', '')
            pdf.multi_cell(0, 4.5, txt=f"  - {bullet_text}")
        else:
            text = cleaned_line.replace('**', '')
            pdf.multi_cell(0, 4.5, txt=text)
            
    pdf.output(filepath)
